import xml.etree.ElementTree as ET
import logging
import docx


logging.basicConfig(level=logging.DEBUG)


def nessus_parse():
    tree = ET.parse('scan.nessus')

    root = tree.getroot()

    vulns = {}

    for item in root:
        logging.debug("\n\nItem: {}-{}-{}".format(item.tag,item.attrib,item.text))
        for thing in item:
            for idk in thing:
                try:
                    if int(idk.attrib["severity"]) == 0:
                        continue
                    else:
                        logging.debug("\nHost: {}".format(thing.attrib["name"]))
                        logging.debug("Plugin Name: {}".format(idk.attrib["pluginName"]))
                        logging.debug("Severity: {}".format(idk.attrib["severity"]))
                        vulns[str(thing.attrib["name"])] = {"Host":thing.attrib["name"],"PluginName":idk.attrib["pluginName"],"Severity":idk.attrib["severity"]}
                        for element in idk:
                            if element.tag == "solution" or element.tag == "see_also" or element.tag == "plugin_output":
                                try:
                                    logging.debug("Tag: {}".format(element.tag))
                                    logging.debug("Text: {}".format(element.text))
                                    logging.debug(thing.attrib["name"])
                                    vulns[str(thing.attrib["name"])].update({element.tag:element.text})
                                except UnicodeEncodeError:
                                    logging.debug("Text: ",element.text)
                except KeyError:
                    continue
    return vulns


def word(vulns):
    doc = docx.Document()
    doc.add_paragraph("Nessus Report for someone")
    doc.add_page_break()
    for x,y in vulns.items():
        doc.add_paragraph(x)
        for z,aa in vulns[x].items():
            doc.add_paragraph("{}\n{}".format(z,aa))
        doc.add_page_break()
    doc.save("hello.docx")
    print("[+] Docx creation completed!")

def main():
    vulns = nessus_parse()
    for x,y in vulns.items():
        print("\n\n\n{}".format(x))
        for z in vulns[x]:
            print("{}\t-\t{}".format(z,vulns[x][z]))

    word(vulns)


if __name__ == "__main__":
    main()
